import tkinter as tk
from tkinter import messagebox, filedialog
from docx import Document
import os

def fill_template(data, template_path, output_path, meetings=None):
    doc = Document(template_path)

    def replace(text):
        for key, value in data.items():
            placeholder = "{{" + key + "}}"
            text = text.replace(placeholder, str(value))
        return text

    if meetings:
        blocks = []
        total_sum = 0
        for meeting in meetings:
            block = (
                f"Дата проведения мероприятия с {meeting['Компания']} {meeting['Дата']}\n"
                f"• организаторов: {meeting['Организаторы']} человека;\n"
                f"• участников: {meeting['Участники']} человек;\n"
                f"• гостей: {meeting['Гости']} человек;\n"
                f"• представительские расходы на сумму {meeting['Сумма']} рублей 00 копеек\n"
            )
            blocks.append(block)
            total_sum += meeting["Сумма"]

        summary = (
            "\nПо итогам представительских мероприятий были приняты решения:\n"
            "• о сотрудничестве с вышеуказанными компаниями в 2025 году\n\n"
            f"Были совершены представительские расходы: на общую сумму {total_sum} рублей 00 копеек.\n"
            "Документы прилагаются.\n"
        )

        data["СписокВстреч"] = "\n".join(blocks) + "\n" + summary
        data["СуммаРасходов"] = total_sum

    for paragraph in doc.paragraphs:
        paragraph.text = replace(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = replace(cell.text)

    doc.save(output_path)

meetings = []

def update_meeting_listbox():
    listbox_meetings.delete(0, tk.END)
    for i, m in enumerate(meetings):
        line = f"{i+1}. {m['Компания']} | {m['Дата']} | орг:{m['Организаторы']} уч:{m['Участники']} гост:{m['Гости']} сумма:{m['Сумма']}"
        listbox_meetings.insert(tk.END, line)

def add_meeting():
    try:
        m = {
            "Компания": entry_meeting_company.get(),
            "Дата": entry_meeting_date.get(),
            "Организаторы": int(entry_meeting_org.get()),
            "Участники": int(entry_meeting_uch.get()),
            "Гости": int(entry_meeting_guest.get()),
            "Сумма": int(entry_meeting_sum.get())
        }
        meetings.append(m)
        update_meeting_listbox()
        clear_meeting_fields()
    except Exception as e:
        messagebox.showerror("Ошибка", str(e))

def clear_meeting_fields():
    entry_meeting_company.delete(0, tk.END)
    entry_meeting_date.delete(0, tk.END)
    entry_meeting_org.delete(0, tk.END)
    entry_meeting_uch.delete(0, tk.END)
    entry_meeting_guest.delete(0, tk.END)
    entry_meeting_sum.delete(0, tk.END)

def clear_form():
    for entry in entries.values():
        entry.delete(0, tk.END)
    meetings.clear()
    update_meeting_listbox()
    clear_meeting_fields()
    messagebox.showinfo("Очищено", "Форма и список встреч очищены.")

def on_select_meeting(event):
    selection = listbox_meetings.curselection()
    if selection:
        idx = selection[0]
        m = meetings[idx]
        entry_meeting_company.delete(0, tk.END)
        entry_meeting_date.delete(0, tk.END)
        entry_meeting_org.delete(0, tk.END)
        entry_meeting_uch.delete(0, tk.END)
        entry_meeting_guest.delete(0, tk.END)
        entry_meeting_sum.delete(0, tk.END)

        entry_meeting_company.insert(0, m["Компания"])
        entry_meeting_date.insert(0, m["Дата"])
        entry_meeting_org.insert(0, m["Организаторы"])
        entry_meeting_uch.insert(0, m["Участники"])
        entry_meeting_guest.insert(0, m["Гости"])
        entry_meeting_sum.insert(0, m["Сумма"])

def update_selected_meeting():
    try:
        idx = listbox_meetings.curselection()[0]
        meetings[idx] = {
            "Компания": entry_meeting_company.get(),
            "Дата": entry_meeting_date.get(),
            "Организаторы": int(entry_meeting_org.get()),
            "Участники": int(entry_meeting_uch.get()),
            "Гости": int(entry_meeting_guest.get()),
            "Сумма": int(entry_meeting_sum.get())
        }
        update_meeting_listbox()
        clear_meeting_fields()
        listbox_meetings.selection_clear(0, tk.END)
    except IndexError:
        messagebox.showwarning("Не выбрано", "Сначала выберите встречу в списке.")

def delete_selected_meeting():
    try:
        idx = listbox_meetings.curselection()[0]
        del meetings[idx]
        update_meeting_listbox()
        clear_meeting_fields()
    except IndexError:
        messagebox.showwarning("Не выбрано", "Сначала выберите встречу в списке.")

def generate_docs():
    try:
        folder = filedialog.askdirectory(title="Выберите папку для сохранения документов")
        if not folder:
            return

        data = {
            "ФИО": entry_fio.get(),
            "ФИОДат": entry_fio_dat.get(),
            "Должность": entry_dolzhnost.get(),
            "Компания": entry_company.get(),
            "НачалоКомандировки": entry_date_start.get(),
            "КонецКомандировки": entry_date_finish.get(),
            "НомерПриказа": entry_order_number.get(),
            "ДатаПриказа": entry_order_date.get(),
            "ДатаОтчета": entry_report_date.get(),
            "Город": entry_city.get(),
            "СуммаПредельныхРасходов": entry_sum_max.get(),
        }

        fio = data["ФИО"]
        start = data["НачалоКомандировки"]
        end = data["КонецКомандировки"]

        fill_template(data, "template_prikaz.docx", os.path.join(folder, f"Приказ_{fio} {start} - {end}.docx"))
        fill_template(data, "template_smeta.docx", os.path.join(folder, f"Смета_{fio} {start} - {end}.docx"))
        fill_template(data, "template_otchet.docx", os.path.join(folder, f"Отчет_{fio} {start} - {end}.docx"), meetings)

        messagebox.showinfo("Успех", f"Документы успешно созданы в папке:\n{folder}")
    except Exception as e:
        messagebox.showerror("Ошибка", str(e))

# Интерфейс
root = tk.Tk()
root.title("Генератор командировочных документов")

fields = [
    ("ФИО", "ФИО"),
    ("ФИО Дат", "ФИОДат"),
    ("Должность", "Должность"),
    ("Компания", "Компания"),
    ("Начало командировки", "НачалоКомандировки"),
    ("Конец командировки", "КонецКомандировки"),
    ("Номер приказа", "НомерПриказа"),
    ("Дата приказа", "ДатаПриказа"),
    ("Дата отчёта", "ДатаОтчета"),
    ("Город", "Город"),
    ("Сумма Предельных Расходов", "СуммаПредельныхРасходов")
]

entries = {}
for label_text, var_name in fields:
    row = tk.Frame(root)
    label = tk.Label(row, width=25, text=label_text + ":", anchor='w')
    entry = tk.Entry(row, width=40)
    row.pack(padx=10, pady=2)
    label.pack(side=tk.LEFT)
    entry.pack(side=tk.RIGHT)
    entries[var_name] = entry

entry_fio = entries["ФИО"]
entry_fio_dat = entries["ФИОДат"]
entry_dolzhnost = entries["Должность"]
entry_company = entries["Компания"]
entry_date_start = entries["НачалоКомандировки"]
entry_date_finish = entries["КонецКомандировки"]
entry_order_number = entries["НомерПриказа"]
entry_order_date = entries["ДатаПриказа"]
entry_report_date = entries["ДатаОтчета"]
entry_city = entries["Город"]
entry_sum_max = entries["СуммаПредельныхРасходов"]

# Блок встреч
tk.Label(root, text="Добавить/Редактировать встречу", font=("Arial", 10, "bold")).pack(pady=5)
meeting_frame = tk.Frame(root)
meeting_frame.pack(pady=3)

entry_meeting_company = tk.Entry(meeting_frame, width=20)
entry_meeting_date = tk.Entry(meeting_frame, width=10)
entry_meeting_org = tk.Entry(meeting_frame, width=5)
entry_meeting_uch = tk.Entry(meeting_frame, width=5)
entry_meeting_guest = tk.Entry(meeting_frame, width=5)
entry_meeting_sum = tk.Entry(meeting_frame, width=10)

tk.Label(meeting_frame, text="Компания").grid(row=0, column=0)
tk.Label(meeting_frame, text="Дата").grid(row=0, column=1)
tk.Label(meeting_frame, text="Организаторы").grid(row=0, column=2)
tk.Label(meeting_frame, text="Участники").grid(row=0, column=3)
tk.Label(meeting_frame, text="Гости").grid(row=0, column=4)
tk.Label(meeting_frame, text="Сумма").grid(row=0, column=5)

entry_meeting_company.grid(row=1, column=0)
entry_meeting_date.grid(row=1, column=1)
entry_meeting_org.grid(row=1, column=2)
entry_meeting_uch.grid(row=1, column=3)
entry_meeting_guest.grid(row=1, column=4)
entry_meeting_sum.grid(row=1, column=5)

# Кнопки встречи
tk.Button(root, text="Добавить встречу", command=add_meeting).pack(pady=2)
tk.Button(root, text="Обновить выбранную встречу", command=update_selected_meeting).pack(pady=2)
tk.Button(root, text="Удалить выбранную встречу", command=delete_selected_meeting).pack(pady=2)

# Список встреч
listbox_meetings = tk.Listbox(root, width=100, height=6)
listbox_meetings.pack(padx=10, pady=5)
listbox_meetings.bind('<<ListboxSelect>>', on_select_meeting)

# Финальные кнопки
tk.Button(root, text="Сгенерировать документы", command=generate_docs).pack(pady=5)
tk.Button(root, text="Очистить форму", command=clear_form).pack(pady=5)

root.mainloop()
